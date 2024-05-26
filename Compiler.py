"""
Fully-automated IAC report compiler
"""
def Compiler(EnergyChartsPath:str, RecommendationPath:str, ReportPath:str):
    import json, os, locale, datetime, math, platform, re, sys
    import pandas as pd
    from easydict import EasyDict
    from docx import Document, shared
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.enum.style import WD_STYLE_TYPE
    from docxcompose.composer import Composer
    from python_docx_replace import docx_replace, docx_blocks
    from Shared.IAC import title_case, dollar, payback, validate_arc, add_image, grouping_num
    import tkinter as tk

    # Check if EnergyChartsPath is empty
    if EnergyChartsPath == "":
        tk.messagebox.showerror("Error", "Energy Charts path is not selected.")
        return
    # Check if RecommendationsPath is empty
    if RecommendationPath == "":
        tk.messagebox.showerror("Error", "Recommendations folder is not selected.")
        return
    # Check if ReportPath is empty
    if ReportPath == "":
        tk.messagebox.showerror("Error", "Report destination path is not selected.")
        return

    # Check if Description.docx has been changed
    docTest = Document(os.path.join('Report', 'Description.docx'))
    # If finds "#Insert plant layout picture here and delete this line", the file has not been changed yet
    for p in docTest.paragraphs:
        if "#Insert plant layout picture here and delete this line" in p.text:
            # pop up a tkinter messagebox
            msb=tk.messagebox.askokcancel("Warning", "You haven't edited Description.docx yet. Continue?")
            if msb == False:
                sys.stderr.write("Please edit Description.docx and compile again.")
                return

    # If the HTML is saved on macOS
    if os.path.exists(os.path.join(EnergyChartsPath, 'EnergyCharts.fld')):
        chartPath = os.path.join(EnergyChartsPath, 'EnergyCharts.fld')
    # If the HTML is saved on Windows
    elif os.path.exists(os.path.join(EnergyChartsPath, 'EnergyCharts_files')):
        chartPath = os.path.join(EnergyChartsPath, 'EnergyCharts_files')
    else:
        # If chart html folder doesn't exist, exit
        tk.messagebox.showerror("Error", "Chart images not found. Please save the energy charts as web page (.htm).")
        return

    # Load json database
    print("Reading json database...", end ="")
    loadJSON = json.load(open('Compiler.json'))
    loadJSON.update(json.load(open('Utility.json')))
    # Create EasyDict
    iac = EasyDict()
    # Convert all sub-dicts to values
    for key in loadJSON:
        iac[key]=loadJSON[key]["value"]
    print("done")

    # Initialize dataframe
    columns = ["isAdditional", "File Name", "ARC No.", "Description", "Electricity (kWh)", "Electricity (MMBtu)", "Demand (kW)"
            , "Natural Gas (MMBtu)", "Other Energy Type", "Other Energy Amount", "Other Resource Type", "Other Resource Amount"
            , "Savings Type", "Savings Value", "Annual Cost Savings", "Implementation Cost", "Payback Period"]
    df = pd.DataFrame(columns=columns)

    # Set locale to en_US
    locale.setlocale(locale.LC_ALL, 'en_US')

    print("Reading recommendations...")
    # Get all .docx files in Recommendations/ directory and extract information
    recList = [f for f in os.listdir(RecommendationPath) if f.endswith('.docx')]
    recID = 0
    for recDoc in recList:
        print(recDoc)
        doc = Document(os.path.join(RecommendationPath, recDoc))
        recInfo = {}
        # Record file name
        recInfo['File Name'] = recDoc

        # Parse document title
        fullTitle = doc.paragraphs[0].text
        separatorFlag = False
        # list of possible separators
        separatorList = [":", "-", "–"]
        for separator in separatorList:
            if separator in fullTitle:
                separatorFlag = True
                # check if the document is an additional recommendation by title
                # Keep "AAR" for outdated documents
                recInfo['isAdditional'] = ("Additional" in fullTitle.split(separator)[0]) or ("AAR" in fullTitle.split(separator)[0])
                # Parse the title of the .docx file
                recInfo['Description'] = title_case(fullTitle.split(separator)[1].strip())
                break
        if separatorFlag == False:
            tk.messagebox.showerror("Error", "Can't parse document title:\n" + fullTitle)
            return
        
        # Read the 1st table in .docx files
        try:
            table = doc.tables[0]
        except:
            tk.messagebox.showerror("Error", recDoc + " is not a valid recommendation. Please check if the summary table is present.")
            return

        for row in table.rows:
            key = row.cells[0].text
            value = row.cells[1].text
            # Parse ARC Number
            if "arc" in key.lower() and "number" in key.lower():
                validate_arc(value)
                recInfo['ARC No.'] = value
            # Parse Annual Cost Savings
            elif "annual" in key.lower() and "cost" in key.lower():
                # convert currency to interger
                recInfo['Annual Cost Savings'] = locale.atoi(value.strip("$"))
            # Parse Implementation Cost
            elif "implementation" in key.lower():
                # convert currency to interger
                recInfo['Implementation Cost'] = locale.atoi(value.strip("$"))
            # If Payback Period skip (Doesn't matter, will calculate later)
            elif "payback" in key.lower():
                continue
            # Parse Electricity
            elif "electricity" in key.lower():
                recInfo['Electricity (kWh)'] = locale.atoi(value.split(' ')[0])
            # Parse Demand
            elif "demand" in key.lower():
                recInfo['Demand (kW)'] = locale.atoi(value.split(' ')[0])
            # Parse Natural Gas
            elif "natural" in key.lower():
                recInfo['Natural Gas (MMBtu)'] = locale.atoi(value.split(' ')[0])
            # Parse undefined type
            else:
                # If the value contains mmbtu, parse it as other energy
                if "mmbtu" in value.lower():
                    # Remove "annual" (usually the first word)
                    if "annual" in key.lower():
                        key = key.split(' ', 1)[1]
                    # Remove "savings" (usually the last word)
                    if "saving" in key.lower():
                        key = key.rsplit(' ', 1)[0]
                    recInfo['Other Energy Type'] = title_case(key)
                    # Parse number
                    recInfo['Other Energy Amount'] = locale.atoi(value.split(' ')[0])
                # If not, parse it as other resource
                else:
                    # Remove "annual" (usually the first word)
                    if "annual" in key.lower():
                        key = key.split(' ', 1)[1]
                    # Remove "savings" (usually the last word)
                    if "saving" in key.lower():
                        key = key.rsplit(' ', 1)[0]
                    recInfo['Other Resource Type'] = title_case(key)
                    # Keep the whole string
                    recInfo['Other Resource Amount'] = value   
        # Add dictionary to dataframe
        for key in recInfo:
            df.loc[recID, key] = recInfo[key]
        recID += 1
    print("done")

    print("Analyzing recommendations...", end ="")
    ## Calculate on columns
    # Calculate payback period
    df['Payback Period'] = df['Implementation Cost'] / df['Annual Cost Savings']
    # Convert electricity to MMBtu
    df['Electricity (MMBtu)'] = df['Electricity (kWh)']* 0.003413/0.33
    # Sort df by payback period
    df = df.sort_values(by=['Payback Period'])

    ## Format energy savings strings
    for index, row in df.iterrows():
        ST = ""
        SV = ""
        if pd.notna(row['Electricity (kWh)']):
            ST = ST + "Electricity" + '\n\n'
            SV = SV + locale.format_string('%d',row['Electricity (kWh)'], grouping=True) + ' kWh' + '\n'
            SV = SV + '(' + locale.format_string('%d',row['Electricity (MMBtu)'], grouping=True) + ' MMBtu)' + '\n'
        if pd.notna(row['Demand (kW)']):
            ST = ST + "Demand" + '\n'
            SV = SV + locale.format_string('%d',row['Demand (kW)'], grouping=True) + ' kW' + '\n'
        if pd.notna(row['Natural Gas (MMBtu)']):
            ST = ST + "Natural Gas" + '\n'
            SV = SV + locale.format_string('%d',row['Natural Gas (MMBtu)'], grouping=True)  + ' MMBtu' + '\n'
        if pd.notna(row['Other Energy Type']):
            ST = ST + row['Other Energy Type'] + '\n'
            SV = SV + locale.format_string('%d',row['Other Energy Amount'], grouping=True)  + ' MMBtu' '\n'
        if pd.notna(row['Other Resource Type']):
            ST = ST + row['Other Resource Type'] + '\n'
            SV = SV + row['Other Resource Amount'] + '\n'
        ST = ST.rstrip('\n')
        SV = SV.rstrip('\n')
        df.at[index, 'Savings Type'] = ST
        df.at[index, 'Savings Value'] = SV

    ## Summation statistics
    # Filter Recommendationss
    recData = df[df['isAdditional'] == False]
    # Reorder index
    recData = recData.reset_index(drop=True)
    # Recommendations statistics
    EkWh = recData['Electricity (kWh)'].sum(axis=0, skipna=True)
    EMMBtu = recData['Electricity (MMBtu)'].sum(axis=0, skipna=True)
    NMMBtu = recData['Natural Gas (MMBtu)'].sum(axis=0, skipna=True)
    OMMBtu = recData['Other Energy Amount'].sum(axis=0, skipna=True)
    # Add up all energy in MMBtu
    iac.MMBtu = round(EMMBtu + NMMBtu + OMMBtu)
    # Calculate CO2
    if iac.FuelType == "Natural Gas":
        iac.FuelCO2 = 53
        iac.CO2 = round((iac.FuelCO2 * NMMBtu + 0.315 * EkWh)/1000)
    elif iac.FuelType == "Propane":
        iac.FuelCO2 = 61.7
        iac.CO2 = round((iac.FuelCO2 * OMMBtu + 0.315 * EkWh)/1000)
    elif iac.FuelType == "Fuel Oil #2":
        iac.FuelCO2 = 73.51
        iac.CO2 = round((iac.FuelCO2 * OMMBtu + 0.315 * EkWh)/1000)
    # Add up all cost
    iac.ACS = recData['Annual Cost Savings'].sum(axis=0, skipna=True)
    iac.IC = recData['Implementation Cost'].sum(axis=0, skipna=True)
    # Payback period in number
    iac.PB = math.ceil(iac.IC / iac.ACS * 10) / 10
    # Payback period in formatted string
    iac.PBstr = payback(iac.ACS, iac.IC)
    print("done")

    print("Reformatting recommendations...", end ="")
    subtitleList = ["Recommended Actions","Summary of Estimated Savings and Implementation Costs","Current Practice and Observations","Anticipated Savings","Implementation Costs","Implementation Cost References"]
    ## Reformatting Recommendations
    for index, row in recData.iterrows():
        doc = Document(os.path.join(RecommendationPath, row['File Name']))
        # Change title and make it upper case
        doc.paragraphs[0].text = "Recommendation "+ str(index+1) + ': ' + title_case(row['Description'])
        # Enforce Heading 1 style
        try:
            doc.paragraphs[0].style = doc.styles['Heading 1']
        except:
            doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            doc.paragraphs[0].style = doc.styles['Heading 1']
        # Enforce Subtitle style
        # This style is already defined in Introduction.docx
        for paragraph in doc.paragraphs:
            txt = paragraph.text
            for subtitle in subtitleList:
                # single or plural
                if txt == subtitle or txt == subtitle[:-1]:
                    try:
                        paragraph.style = doc.styles['Subtitle']
                    except:
                        doc.styles.add_style('Subtitle', WD_STYLE_TYPE.PARAGRAPH)
                        paragraph.style = doc.styles['Subtitle']
            # Fix table/figure captions
            if re.search('^\s?Table\s\d{1,2}:', txt) != None or re.search('^\s?Figure\s\d{1,2}:', txt) != None:
                try:
                    paragraph.style = doc.styles['Caption']
                except:
                    doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
                    paragraph.style = doc.styles['Caption']

        # Save file with sorted filename
        doc.save(os.path.join(RecommendationPath, 'Sorted', 'Rec'+ str(index+1) + '.docx'))
    print("done")

    # Check if there's at least 1 additional recommendation
    hasAdditional = df['isAdditional'].any()
    if hasAdditional:
        print("Analyzing additional recommendations...", end ="")
        # Filter additional
        addData = df[df['isAdditional'] == True]
        # Reorder index
        addData = addData.reset_index(drop=True)
        # Additional statistics
        EMMBtu = addData['Electricity (MMBtu)'].sum(axis=0, skipna=True)
        NMMBtu = addData['Natural Gas (MMBtu)'].sum(axis=0, skipna=True)
        OMMBtu = addData['Other Energy Amount'].sum(axis=0, skipna=True)
        # Add up all energy
        iac.AddMMBtu = round(EMMBtu + NMMBtu + OMMBtu)
        # Add up all cost
        iac.AddACS = addData['Annual Cost Savings'].sum(axis=0, skipna=True)
        iac.AddIC = addData['Implementation Cost'].sum(axis=0, skipna=True)
        # Payback period in number
        iac.AddPB = round(iac.AddIC / iac.AddACS, 1)
        print("done")

        print("Reformatting additional recommendations...", end ="")
        # Modify the title of the additional recommendation docx
        for index, row in addData.iterrows():
            doc = Document(os.path.join(RecommendationPath, row['File Name']))
            # Change title and make it upper case
            doc.paragraphs[0].text = "Additional Recommendation "+ str(index+1) + ': ' + title_case(row['Description'])
            # Enforce Heading 1 style
            try:
                doc.paragraphs[0].style = doc.styles['Heading 1']
            except:
                doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
                doc.paragraphs[0].style = doc.styles['Heading 1']
            # Enforce Subtitle style
            # This style is already defined in Introduction.docx
            for paragraph in doc.paragraphs:
                txt = paragraph.text
                for subtitle in subtitleList:
                    # single or plural
                    if txt == subtitle or txt == subtitle[:-1]:
                        try:
                            paragraph.style = doc.styles['Subtitle']
                        except:
                            doc.styles.add_style('Subtitle', WD_STYLE_TYPE.PARAGRAPH)
                            paragraph.style = doc.styles['Subtitle']
                # Fix table/figure captions
                if re.search('^\s?Table\s\d{1,2}:', txt) != None or re.search('^\s?Figure\s\d{1,2}:', txt) != None:
                    try:
                        paragraph.style = doc.styles['Caption']
                    except:
                        doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
                        paragraph.style = doc.styles['Caption']
            # Save file with sorted filename
            doc.save(os.path.join(RecommendationPath, 'Sorted', 'Add'+ str(index+1) + '.docx'))
        print("done")

    print("Parsing plant information...", end ="")

    ## Compiler.json Calculations
    # Report date = today or 60 days after assessment, which ever is earlier
    VD = datetime.datetime.strptime(iac.VDATE, '%m/%d/%Y')
    RDATE = min(datetime.datetime.today(), VD + datetime.timedelta(days=60))
    if platform.system() == 'Windows':
        iac.VDATE = datetime.datetime.strftime(VD, '%B %#d, %Y')
        iac.RDATE = datetime.datetime.strftime(RDATE, '%B %#d, %Y')
    else: # macOS or Linux
        iac.VDATE = datetime.datetime.strftime(VD, '%B %-d, %Y')
        iac.RDATE = datetime.datetime.strftime(RDATE, '%B %-d, %Y')

    # Format operating hours
    PROH = iac.PROH
    iac.PROH = str(PROH[0]) + " hours per day, " + str(PROH[1]) + " days per week, " + str(PROH[2]) + " weeks per year"
    OFOH = iac.OFOH
    iac.OFOH = str(OFOH[0]) + " hours per day, " + str(OFOH[1]) + " days per week, " + str(OFOH[2]) + " weeks per year"

    # Sort participant and contributor name list
    PARTlist=iac.PART
    PARTlist.sort(key=lambda x: x.rsplit(' ', 1)[1])
    PART=""
    for name in PARTlist:
        PART  = PART + name + '\n'
    iac.PART = PART.rstrip('\n')

    CONTlist=iac.CONT
    CONTlist.sort(key=lambda x: x.rsplit(' ', 1)[1])
    CONT=""
    for name in CONTlist:
        CONT  = CONT + name + '\n'
    iac.CONT = CONT.rstrip('\n')
    print("done")

    # products in different cases
    iac.PRODTitle = iac.PROD.title()
    iac.PRODlower = iac.PROD.lower()

    ## Format strings
    # set electricity cost to 3 digits accuracy
    iac = dollar(['EC'],iac,3)
    # set the natural gas and demand to 2 digits accuracy
    iac = dollar(['DC', 'FC'],iac,2)
    # set the rest to integer
    varList = ['ACS', 'IC', 'TotalECost', 'TotalFCost', 'TotalCost']
    if hasAdditional:
        varList.extend(['AddACS', 'AddIC'])
    iac = dollar(varList,iac,0)
    # Format all numbers to string with thousand separator
    iac = grouping_num(iac)

    ## Load introduction template
    docIntro = Document(os.path.join('Report', 'Introduction.docx'))

    # Add rows to Recommendation table (Should be the 3rd table)
    print("Writing recommendation table...", end ="")
    locale._override_localeconv={'frac_digits':0}
    recTable = docIntro.tables[2]
    for index, row in recData.iterrows():
        recRow = recTable.rows[index+1].cells
        # Add ARC No.
        recRow[0].text = 'Rec. ' + str(index+1) + '\n' + row['ARC No.']
        recRow[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add description
        recRow[1].text = row['Description']
        recRow[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Add savings type
        recRow[2].text = row['Savings Type']
        recRow[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add savings value
        recRow[3].text = row['Savings Value']
        recRow[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add annual cost savings
        recRow[4].text = locale.currency(row['Annual Cost Savings'], grouping=True)
        recRow[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Add implementation cost
        recRow[5].text = locale.currency(row['Implementation Cost'], grouping=True)
        recRow[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Add payback period
        pb = row['Payback Period']
        if pb == 0:
            recRow[6].text = "Immediate"
        else:
            recRow[6].text = str(math.ceil(pb * 10) / 10)
        recRow[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Set 3pt before and after paragraph
        for col in range(0,7):
            recRow[col].paragraphs[0].paragraph_format.space_before = shared.Pt(3)
            recRow[col].paragraphs[0].paragraph_format.space_after = shared.Pt(3)
    # Delete unused rows (Currectly row 1-15 are empty)
    for index in reversed(range(len(recData), 15)):
        recTable._tbl.remove(recTable.rows[index+1]._tr)
    print("done")

    if hasAdditional:
        # Add rows to additional recommendation table (Should be the 4th table)
        print("Writing Additional Recommendation table...", end ="")
        locale._override_localeconv={'frac_digits':0}
        addTable = docIntro.tables[3]
        for index, row in addData.iterrows():
            addRow = addTable.rows[index+1].cells
            # Add ARC No.
            addRow[0].text = 'Add. Rec. ' + str(index+1) + '\n' + row['ARC No.']
            addRow[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add description
            addRow[1].text = row['Description']
            addRow[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Add savings type
            addRow[2].text = row['Savings Type']
            addRow[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add savings value
            addRow[3].text = row['Savings Value']
            addRow[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add annual cost savings
            addRow[4].text = locale.currency(row['Annual Cost Savings'], grouping=True)
            addRow[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Add implementation cost
            addRow[5].text = locale.currency(row['Implementation Cost'], grouping=True)
            addRow[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Add payback period
            pb = row['Payback Period']
            addRow[6].text = str(math.ceil(pb * 10) / 10)
            addRow[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Set 3pt before and after paragraph
        for col in range(0,7):
            addRow[col].paragraphs[0].paragraph_format.space_before = shared.Pt(3)
            addRow[col].paragraphs[0].paragraph_format.space_after = shared.Pt(3)
        # Delete unused rows (Currectly row 1-5 are empty)
        for index in reversed(range(len(addData), 5)):
            addTable._tbl.remove(addTable.rows[index+1]._tr)
        print("done")
    else:
        # delete this table
        docIntro._body._body.remove(docIntro.tables[3]._tbl)

    # Remove Add blocks if no Additional
    docx_blocks(docIntro, ADD = hasAdditional)

    # Replacing keys
    print("Replacing keys in introduction...", end ="")
    docx_replace(docIntro, **iac)
    print("done")

    # Save introduction
    filenameIntro = iac.LE + '-intro.docx'
    docIntro.save(filenameIntro)

    ## Load backgroud template
    docBackground = Document(os.path.join('Report', 'Background.docx'))

    # Replacing keys
    print("Replacing keys in background...", end ="")
    docx_replace(docBackground, **iac)
    print("done")

    # Save background
    filenameBackground = iac.LE + '-back.docx'
    docBackground.save(filenameBackground)

    ## Load energy bill analysis template
    docEnergy = Document(os.path.join('Report', 'Energy.docx'))

    # Add energy charts images
    print("Adding energy charts images...", end ="")
    # If on macOS
    if chartPath == os.path.join(EnergyChartsPath, 'EnergyCharts.fld'):
        add_image(docEnergy, '#EUChart', os.path.join(chartPath, "image001.png"), shared.Inches(6))
        add_image(docEnergy, '#ECChart', os.path.join(chartPath, "image004.png"), shared.Inches(6))
        add_image(docEnergy, '#DUChart', os.path.join(chartPath, "image002.png"), shared.Inches(6))
        add_image(docEnergy, '#DCChart', os.path.join(chartPath, "image005.png"), shared.Inches(6))
        add_image(docEnergy, '#FUChart', os.path.join(chartPath, "image003.png"), shared.Inches(6))
        add_image(docEnergy, '#FCChart', os.path.join(chartPath, "image006.png"), shared.Inches(6))
        add_image(docEnergy, '#PieUChart', os.path.join(chartPath, "image007.png"), shared.Inches(6))
        add_image(docEnergy, '#PieCChart', os.path.join(chartPath, "image008.png"), shared.Inches(6))
        add_image(docEnergy, '#TotalChart', os.path.join(chartPath, "image009.png"), shared.Inches(9))
    # If on Windows
    elif chartPath == os.path.join(EnergyChartsPath, 'EnergyCharts_files'):
        add_image(docEnergy, '#EUChart', os.path.join(chartPath, "image001.png"), shared.Inches(6))
        add_image(docEnergy, '#ECChart', os.path.join(chartPath, "image002.png"), shared.Inches(6))
        add_image(docEnergy, '#DUChart', os.path.join(chartPath, "image003.png"), shared.Inches(6))
        add_image(docEnergy, '#DCChart', os.path.join(chartPath, "image005.png"), shared.Inches(6))
        add_image(docEnergy, '#FUChart', os.path.join(chartPath, "image006.png"), shared.Inches(6))
        add_image(docEnergy, '#FCChart', os.path.join(chartPath, "image007.png"), shared.Inches(6))
        add_image(docEnergy, '#PieUChart', os.path.join(chartPath, "image009.png"), shared.Inches(6))
        add_image(docEnergy, '#PieCChart', os.path.join(chartPath, "image011.png"), shared.Inches(6))
        add_image(docEnergy, '#TotalChart', os.path.join(chartPath, "image013.png"), shared.Inches(9))
    print("done")

    # Fill in energy charts tables from EnergyCharts.xlsx
    print("Adding energy charts tables...", end ="")
    # Read electricity table from B6 to I19
    edf = pd.read_excel(os.path.join(EnergyChartsPath, 'EnergyCharts.xlsx'), sheet_name="Raw Data", skiprows = 5, nrows=13, usecols = 'B:I')
    # Read fuel table from K6 to N19
    fdf = pd.read_excel(os.path.join(EnergyChartsPath, 'EnergyCharts.xlsx'), sheet_name="Raw Data", skiprows = 5, nrows=13, usecols = 'K:N')

    # Add rows to electricity table (Should be the 1st table)
    eTable = docEnergy.tables[0]
    for index, row in edf.iterrows():
        eRow = eTable.rows[index+3].cells
        # Add Month
        eRow[0].text = edf.iloc[(index, 0)]
        eRow[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for col in range(1,8):
            # Add interger with thousand separator
            eRow[col].text = locale.format_string('%d',round(edf.iloc[(index, col)]), grouping=True)
            eRow[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Bold the last row
        if index == 12:
            for col in range(0,8):
                eRow[col].paragraphs[0].runs[0].bold = True

    # Add rows to fuel table (Should be the 2nd table)
    fTable = docEnergy.tables[1]
    for index, row in fdf.iterrows():
        fRow = fTable.rows[index+3].cells
        # Add Month
        fRow[0].text = fdf.iloc[(index, 0)]
        fRow[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for col in range(1,4):
            # Add interger with thousand separator
            fRow[col].text = locale.format_string('%d',round(fdf.iloc[(index, col)]), grouping=True)
            fRow[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Bold the last row
        if index == 12:
            for col in range(0,4):
                fRow[col].paragraphs[0].runs[0].bold = True
    print("done")
    # Replacing keys
    print("Replacing keys in energycharts...", end ="")
    docx_replace(docEnergy, **iac)
    print("done")
    # Save energyc harts docx
    filenameEnergy = iac.LE + '-energy.docx'
    docEnergy.save(filenameEnergy)

    print("Combining all docs...", end ="")
    # List of docs to combine
    docList = [os.path.join('Report', 'ToC.docx')]
    for RecLength in range(1, len(recData)+1):
        docList.append(os.path.join(RecommendationPath, 'Sorted','Rec' + str(RecLength) + '.docx'))
    if hasAdditional:
        docList.append(os.path.join('Report', 'Add.docx'))
        for AddLength in range(1, len(addData)+1):
            docList.append(os.path.join(RecommendationPath, 'Sorted','Add' + str(AddLength) + '.docx'))
    else:
        pass
    docList.append(filenameBackground)
    docList.append(os.path.join('Report', 'Description.docx'))

    # Combine all docx files
    main = Document(filenameIntro)
    main.add_page_break()
    composer = Composer(main)
    for i in range(0, len(docList)):
        doc_add = Document(docList[i])
        doc_add.add_page_break()
        composer.append(doc_add)
    # A section break is already added in BestPractice.docx, so no need to add a page break
    composer.append(Document(os.path.join('Report', 'BestPractice.docx')))
    composer.append(Document(filenameEnergy))
    composer.save(ReportPath)

    # delete temp files
    os.remove(filenameIntro)
    os.remove(filenameBackground)
    os.remove(filenameEnergy)
    print("done")

    # Open the combined docx file
    doc = Document(ReportPath)
    # Change the orientation of the last section to landscape
    section = doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    # Save final report
    doc.save(ReportPath)
    print(ReportPath + " is finished.")